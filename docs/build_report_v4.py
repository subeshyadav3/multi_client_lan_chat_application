#!/usr/bin/env python3
"""
Build final_report_v4.docx following Tribhuvan University (IOE Pulchowk Campus)
final project report standards. 
- Pure black typography (all blue removed)
- 1.5 line spacing and proper paragraph spacing
- Each main chapter/topic starts on a NEW PAGE
- Header text removed
- Fixed single footer page numbering (no duplicated numbers)
- Updated client names: Prabesh, Saroj, Subesh
"""

import re
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE  = os.path.join(DOCS_DIR, "final_report_v3.md")
OUT_FILE = os.path.join(DOCS_DIR, "final_report_v4.docx")

BLACK = RGBColor(0x00, 0x00, 0x00)

# ── XML & Formatting Helpers ──────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_section_page_numbering(section, fmt='decimal', start=None):
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

def add_page_number_to_paragraph(p):
    p.text = "" # Clear paragraph runs
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_toc_line(doc, number, title, page_str, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = 1.3

    indent = (level - 1) * 0.25
    p.paragraph_format.left_indent = Inches(indent)

    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '8640') # 6.0 inches right alignment tab stop
    tabs.append(tab)
    pPr.append(tabs)

    if number:
        run_num = p.add_run(f"{number} ")
        run_num.bold = (level == 1)
        run_num.font.name = "Times New Roman"
        run_num.font.size = Pt(11)
        run_num.font.color.rgb = BLACK

    run_title = p.add_run(title)
    run_title.bold = (level == 1)
    run_title.font.name = "Times New Roman"
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = BLACK

    run_tab = p.add_run(f"\t{page_str}")
    run_tab.bold = (level == 1)
    run_tab.font.name = "Times New Roman"
    run_tab.font.size = Pt(11)
    run_tab.font.color.rgb = BLACK

def style_table(tbl):
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(tbl.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(r'<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
        if i == 0:
            trPr.append(parse_xml(r'<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
        for cell in row.cells:
            if i == 0:
                set_cell_bg(cell, "333333") # Dark Gray Header
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(10)
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                if i % 2 == 0:
                    set_cell_bg(cell, "F9FAFB")
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(10)
                        r.font.color.rgb = BLACK

def embed_image(doc, img_path, caption=None, width=5.5):
    full = os.path.join(DOCS_DIR, img_path)
    if not os.path.exists(full):
        print(f"  [WARN] image not found: {full}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run()
    run.add_picture(full, width=Inches(width))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after  = Pt(14)
        c_run = cp.add_run(caption)
        c_run.font.name = "Times New Roman"
        c_run.font.size = Pt(10)
        c_run.italic = True
        c_run.bold = True
        c_run.font.color.rgb = BLACK

# ── Main Document Construction ────────────────────────────────────────────────

def build():
    doc = Document()

    # Configure default styles
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = BLACK

    # -------------------------------------------------------------------------
    # SECTION 0: COVER PAGE (TU IOE Pulchowk Format)
    # -------------------------------------------------------------------------
    sec_cover = doc.sections[0]
    sec_cover.top_margin    = Inches(1.0)
    sec_cover.bottom_margin = Inches(1.0)
    sec_cover.left_margin   = Inches(1.25) # Binding allowance
    sec_cover.right_margin  = Inches(1.0)
    sec_cover.different_first_page_header_footer = True

    # Header / University
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TRIBHUVAN UNIVERSITY\nINSTITUTE OF ENGINEERING\nPULCHOWK CAMPUS")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"
    r.font.color.rgb = BLACK

    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(36)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONNECTHUB: A MULTI-CLIENT LAN CHAT AND FILE SHARING APPLICATION")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Times New Roman"
    r.font.color.rgb = BLACK

    p_sp2 = doc.add_paragraph()
    p_sp2.paragraph_format.space_before = Pt(36)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run("A PROJECT REPORT\nSUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR THE DEGREE OF BACHELOR OF ENGINEERING IN ELECTRONICS AND COMPUTER ENGINEERING")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"
    r.font.color.rgb = BLACK

    p_sp3 = doc.add_paragraph()
    p_sp3.paragraph_format.space_before = Pt(48)

    # Table for Submitted By / Submitted To
    tbl_cov = doc.add_table(rows=1, cols=2)
    tbl_cov.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_by = tbl_cov.cell(0, 0)
    cell_to = tbl_cov.cell(0, 1)

    p_by = cell_by.paragraphs[0]
    r_by_h = p_by.add_run("Submitted by:\n")
    r_by_h.bold = True
    r_by_h.font.size = Pt(11)
    r_by_h.font.color.rgb = BLACK
    r_by_body = p_by.add_run("Prabesh BC (080BCT054)\nSaroj Rawal (080BCT076)\nSubesh Yadav (080BCT084)")
    r_by_body.font.size = Pt(11)
    r_by_body.font.color.rgb = BLACK

    p_to = cell_to.paragraphs[0]
    p_to.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_to_h = p_to.add_run("Submitted to:\n")
    r_to_h.bold = True
    r_to_h.font.size = Pt(11)
    r_to_h.font.color.rgb = BLACK
    r_to_body = p_to.add_run("Department of Electronics &\nComputer Engineering\nPulchowk Campus, IOE\nTribhuvan University, Nepal")
    r_to_body.font.size = Pt(11)
    r_to_body.font.color.rgb = BLACK

    p_sp4 = doc.add_paragraph()
    p_sp4.paragraph_format.space_before = Pt(64)

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run("August, 2026")
    r_date.bold = True
    r_date.font.size = Pt(12)
    r_date.font.name = "Times New Roman"
    r_date.font.color.rgb = BLACK

    # -------------------------------------------------------------------------
    # SECTION 1: FRONT MATTER (Roman Page Numbers i, ii, iii...)
    # -------------------------------------------------------------------------
    sec_front = doc.add_section(WD_SECTION.NEW_PAGE)
    sec_front.top_margin    = Inches(1.0)
    sec_front.bottom_margin = Inches(1.0)
    sec_front.left_margin   = Inches(1.25)
    sec_front.right_margin  = Inches(1.0)
    sec_front.header.is_linked_to_previous = False
    sec_front.footer.is_linked_to_previous = False
    set_section_page_numbering(sec_front, fmt="lowerRoman", start=1)

    # Footer Roman page numbers
    add_page_number_to_paragraph(sec_front.footer.paragraphs[0])

    # --- Abstract ---
    h_abs = doc.add_heading("Abstract", level=1)
    h_abs.paragraph_format.space_before = Pt(0)
    h_abs.paragraph_format.space_after  = Pt(12)
    for r in h_abs.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = BLACK

    p_abs = doc.add_paragraph(
        "ConnectHub is a working multi-client LAN chat and file-sharing application written entirely in plain C. "
        "It uses nothing beyond the C standard library and standard POSIX calls — no external frameworks, no third-party libraries. "
        "Every part of the system, from the network layer to the terminal interface, is built from scratch using system calls taught in the "
        "undergraduate curriculum: socket, bind, listen, accept, pthread_create, select, read, and write.\n\n"
        "The application follows a classic client–server design. A single multithreaded server manages all connected users, chat rooms, "
        "and file transfers. Each client gets its own POSIX thread on the server, and shared data is protected by mutexes. On the client side, "
        "a single select() loop watches both the keyboard and the network at the same time, so the terminal stays responsive even during a file upload.\n\n"
        "All messages follow a custom text-based protocol — pipe-delimited, newline-terminated lines — that is easy to read and debug with any terminal tool. "
        "The system supports room-based chat, private messaging, typing indicators, join/leave notifications, room history, SHA-256 password authentication, "
        "and bidirectional file transfer with token-based access control.\n\n"
        "Keywords: TCP sockets, POSIX threads, select(), concurrent server, LAN chat, file transfer, custom protocol, C programming, SHA-256"
    )
    p_abs.paragraph_format.line_spacing = 1.5
    p_abs.paragraph_format.space_after  = Pt(12)
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # --- Acknowledgements ---
    h_ack = doc.add_heading("Acknowledgements", level=1)
    h_ack.paragraph_format.space_before = Pt(0)
    h_ack.paragraph_format.space_after  = Pt(12)
    for r in h_ack.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = BLACK

    p_ack = doc.add_paragraph(
        "We thank the Department of Electronics and Computer Engineering, Pulchowk Campus, for giving us the opportunity to undertake this project as part of our Bachelor's program.\n\n"
        "We are grateful to our project supervisor for the guidance and feedback provided throughout the design and implementation phases. The advice received at key decision points significantly shaped the final outcome.\n\n"
        "We also acknowledge the faculty members whose lectures on systems programming, computer networking, and operating systems laid the conceptual foundation for this work. The understanding we gained from those courses directly informed every design choice in ConnectHub.\n\n"
        "Finally, we thank the authors of the references listed at the end of this report, whose published work served as an essential guide."
    )
    p_ack.paragraph_format.line_spacing = 1.5
    p_ack.paragraph_format.space_after  = Pt(12)
    p_ack.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # --- Table of Contents ---
    h_toc = doc.add_heading("Table of Contents", level=1)
    h_toc.paragraph_format.space_before = Pt(0)
    h_toc.paragraph_format.space_after  = Pt(12)
    for r in h_toc.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = BLACK

    # Word XML TOC Field
    p_field = doc.add_paragraph()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    p_field._p.append(fldSimple)

    # Formatted Dot-Leader Entries for viewing on any reader
    toc_data = [
        ("1.", "Introduction", "1", 1),
        ("1.1", "Background", "1", 2),
        ("1.2", "Motivation", "1", 2),
        ("1.3", "Scope", "2", 2),
        ("2.", "Problem Statement", "3", 1),
        ("3.", "Objectives", "4", 1),
        ("3.1", "Main Objective", "4", 2),
        ("3.2", "Specific Objectives", "4", 2),
        ("4.", "Literature Review", "5", 1),
        ("4.1", "Theoretical Background", "5", 2),
        ("4.2", "Related Works", "5", 2),
        ("4.3", "Gap Analysis", "6", 2),
        ("5.", "System Design and Architecture", "7", 1),
        ("5.1", "System Overview", "7", 2),
        ("5.2", "Server Architecture", "8", 2),
        ("5.3", "Client Architecture", "9", 2),
        ("5.4", "Wire Protocol Design", "9", 2),
        ("5.5", "Data Structures", "10", 2),
        ("5.6", "Module Structure", "11", 2),
        ("6.", "Implementation", "12", 1),
        ("6.1", "Technology Stack", "12", 2),
        ("6.2", "Module Breakdown", "12", 2),
        ("6.3", "Authentication", "13", 2),
        ("6.4", "Room Management", "13", 2),
        ("6.5", "Messaging", "14", 2),
        ("6.6", "File Transfer", "14", 2),
        ("6.7", "Build System", "15", 2),
        ("7.", "Testing and Validation", "16", 1),
        ("7.1", "Integration Testing", "16", 2),
        ("7.2", "Concurrency and Stress Testing", "16", 2),
        ("7.3", "Wire-Level Verification", "17", 2),
        ("8.", "Results and Discussion", "18", 1),
        ("9.", "Conclusion and Future Work", "19", 1),
        ("9.1", "Conclusion", "19", 2),
        ("9.2", "Future Work", "19", 2),
        ("10.", "References", "20", 1),
    ]
    for num, title, pg, lvl in toc_data:
        add_toc_line(doc, num, title, pg, level=lvl)

    doc.add_page_break()

    # --- List of Figures ---
    h_lof = doc.add_heading("List of Figures", level=1)
    h_lof.paragraph_format.space_before = Pt(0)
    h_lof.paragraph_format.space_after  = Pt(12)
    for r in h_lof.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = BLACK

    lof_data = [
        ("Figure 1", "LAN Deployment Topology", "7"),
        ("Figure 2", "Two-Tier Client–Server Architecture", "7"),
        ("Figure 3", "Server Thread Model and Locking Strategy", "8"),
        ("Figure 4", "File Transfer Sequence", "14"),
        ("Figure 5", "Module Dependency Map", "11"),
    ]
    for num, title, pg in lof_data:
        add_toc_line(doc, num, title, pg, level=1)

    doc.add_page_break()

    # --- List of Tables ---
    h_lot = doc.add_heading("List of Tables", level=1)
    h_lot.paragraph_format.space_before = Pt(0)
    h_lot.paragraph_format.space_after  = Pt(12)
    for r in h_lot.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = BLACK

    lot_data = [
        ("Table 1", "Gap Analysis Summary", "6"),
        ("Table 2", "Client → Server Protocol Messages", "9"),
        ("Table 3", "Server → Client Protocol Messages", "10"),
        ("Table 4", "Server Shared Data Structures", "10"),
        ("Table 5", "Codebase Module Breakdown", "11"),
        ("Table 6", "ConnectHub Technology Stack", "12"),
        ("Table 7", "Concurrency and Stress Test Results", "16"),
        ("Table 8", "Project Objectives Outcome", "18"),
    ]
    for num, title, pg in lot_data:
        add_toc_line(doc, num, title, pg, level=1)

    doc.add_page_break()

    # --- List of Abbreviations ---
    h_loa = doc.add_heading("List of Abbreviations", level=1)
    h_loa.paragraph_format.space_before = Pt(0)
    h_loa.paragraph_format.space_after  = Pt(12)
    for r in h_loa.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = BLACK

    abbs = [
        ("API", "Application Programming Interface"),
        ("CLI", "Command-Line Interface"),
        ("DM", "Direct Message"),
        ("FD", "File Descriptor"),
        ("I/O", "Input/Output"),
        ("IRC", "Internet Relay Chat"),
        ("LAN", "Local Area Network"),
        ("POSIX", "Portable Operating System Interface"),
        ("SHA", "Secure Hash Algorithm"),
        ("TCP", "Transmission Control Protocol"),
        ("TUI", "Terminal User Interface"),
        ("TU", "Tribhuvan University"),
        ("IOE", "Institute of Engineering"),
    ]
    tbl_abb = doc.add_table(rows=len(abbs)+1, cols=2)
    tbl_abb.cell(0, 0).paragraphs[0].text = "Abbreviation"
    tbl_abb.cell(0, 1).paragraphs[0].text = "Meaning"
    for idx, (a, m) in enumerate(abbs):
        tbl_abb.cell(idx+1, 0).paragraphs[0].text = a
        tbl_abb.cell(idx+1, 1).paragraphs[0].text = m
    style_table(tbl_abb)

    # -------------------------------------------------------------------------
    # SECTION 2: MAIN CONTENT (Arabic Page Numbers 1, 2, 3...)
    # -------------------------------------------------------------------------
    sec_main = doc.add_section(WD_SECTION.NEW_PAGE)
    sec_main.top_margin    = Inches(1.0)
    sec_main.bottom_margin = Inches(1.0)
    sec_main.left_margin   = Inches(1.25)
    sec_main.right_margin  = Inches(1.0)
    sec_main.header.is_linked_to_previous = False
    sec_main.footer.is_linked_to_previous = False
    set_section_page_numbering(sec_main, fmt="decimal", start=1)

    # Footer Arabic page numbers
    add_page_number_to_paragraph(sec_main.footer.paragraphs[0])

    # --- Parse MD content ---
    with open(MD_FILE, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_code_block  = False
    in_table       = False
    table_rows     = []
    code_lines     = []
    skip_front     = True
    first_chapter  = True

    def flush_table():
        nonlocal table_rows, in_table
        if table_rows:
            clean = [r for r in table_rows if not all(re.match(r'^[-| ]+$', c) for c in r)]
            if clean:
                cols = len(clean[0])
                tbl = doc.add_table(rows=0, cols=cols)
                for row_data in clean:
                    tr = tbl.add_row()
                    for j, cell_text in enumerate(row_data):
                        tr.cells[j].paragraphs[0].text = cell_text.strip()
                style_table(tbl)
                doc.add_paragraph()
        table_rows = []
        in_table   = False

    def flush_code():
        nonlocal code_lines, in_code_block
        code_text = "".join(code_lines).rstrip("\n")
        if code_text:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = BLACK
        code_lines   = []
        in_code_block = False

    while i < len(lines):
        line = lines[i]
        raw  = line.rstrip("\n")

        if skip_front:
            if raw.startswith("## 1. Introduction"):
                skip_front = False
            else:
                i += 1
                continue

        # Code block
        if raw.startswith("```"):
            if in_code_block:
                flush_code()
            else:
                if in_table: flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table row
        if raw.startswith("|"):
            cells = [c for c in raw.split("|")]
            cells = cells[1:-1] if cells[0] == "" and cells[-1] == "" else cells
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table: flush_table()

        # Image
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', raw)
        if img_match:
            caption  = img_match.group(1)
            img_path = img_match.group(2)
            embed_image(doc, img_path, caption=caption)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', raw.strip()):
            i += 1
            continue

        # Headings
        h_match = re.match(r'^(#{1,4})\s+(.*)', raw)
        if h_match:
            level = len(h_match.group(1))
            text  = h_match.group(2).strip()
            text  = re.sub(r'`([^`]+)`', r'\1', text)
            
            # Each main Chapter topic (level 1 heading) starts on a NEW PAGE!
            if level == 1:
                if not first_chapter:
                    doc.add_page_break()
                else:
                    first_chapter = False

            p = doc.add_heading(text, level=level)
            p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
            p.paragraph_format.space_after  = Pt(6)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.bold = True
                r.font.color.rgb = BLACK
                if level == 1: r.font.size = Pt(16)
                elif level == 2: r.font.size = Pt(14)
                else: r.font.size = Pt(12)
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r'^(\s*)[*\-]\s+(.*)', raw)
        if bullet_match:
            text = bullet_match.group(2).strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.5
            segments = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
            for seg in segments:
                if seg.startswith("**") and seg.endswith("**"):
                    r = p.add_run(seg[2:-2])
                    r.bold = True
                    r.font.name = "Times New Roman"
                    r.font.color.rgb = BLACK
                elif seg.startswith("`") and seg.endswith("`"):
                    r = p.add_run(seg[1:-1])
                    r.font.name = "Courier New"
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = BLACK
                else:
                    r = p.add_run(seg)
                    r.font.name = "Times New Roman"
                    r.font.color.rgb = BLACK
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^\d+\.\s+(.*)', raw)
        if num_match:
            text = num_match.group(1).strip()
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.5
            segments = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
            for seg in segments:
                if seg.startswith("**") and seg.endswith("**"):
                    r = p.add_run(seg[2:-2])
                    r.bold = True
                    r.font.name = "Times New Roman"
                    r.font.color.rgb = BLACK
                elif seg.startswith("`") and seg.endswith("`"):
                    r = p.add_run(seg[1:-1])
                    r.font.name = "Courier New"
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = BLACK
                else:
                    r = p.add_run(seg)
                    r.font.name = "Times New Roman"
                    r.font.color.rgb = BLACK
            i += 1
            continue

        # Blank line
        if raw.strip() == "":
            i += 1
            continue

        # Normal paragraph
        text = raw.strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after  = Pt(6)

        segments = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
        for seg in segments:
            if seg.startswith("**") and seg.endswith("**"):
                run = p.add_run(seg[2:-2])
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.color.rgb = BLACK
            elif seg.startswith("`") and seg.endswith("`"):
                run = p.add_run(seg[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(9.5)
                run.font.color.rgb = BLACK
            else:
                run = p.add_run(seg)
                run.font.name = "Times New Roman"
                run.font.color.rgb = BLACK

        i += 1

    if in_code_block: flush_code()
    if in_table: flush_table()

    doc.save(OUT_FILE)
    print(f"Saved v4 report to: {OUT_FILE}")

if __name__ == "__main__":
    build()
