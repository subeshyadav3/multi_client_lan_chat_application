import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_docx(filename):
    doc = docx.Document()
    
    # Page setup - Standard 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # Color Palette
    PRIMARY = RGBColor(0x1B, 0x36, 0x5D)    # Deep Navy
    SECONDARY = RGBColor(0x00, 0x66, 0x99)  # Slate Blue
    ACCENT = RGBColor(0xD9, 0x6B, 0x27)     # Amber/Orange
    DARK_TEXT = RGBColor(0x22, 0x22, 0x22)
    MUTED = RGBColor(0x55, 0x55, 0x55)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("ConnectHub: Complete Defense & Viva Master Guide")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    run_sub = p_sub.add_run("Systems Programming, Sockets, POSIX Concurrency, Wire Protocol & Live Demo Walkthrough\nAuthor: Subesh Yadav (080BCT084) | Pulchowk Campus, IOE, Tribhuvan University")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = MUTED
    
    # Horizontal rule
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_after = Pt(14)
    r_hr = p_hr.add_run("―" * 55)
    r_hr.font.color.rgb = SECONDARY
    
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = PRIMARY
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = SECONDARY
        return h

    def add_h3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(2)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11.5)
        r.font.bold = True
        r.font.color.rgb = ACCENT
        return h

    def add_callout(title, text, hex_bg="F0F4F8", border_hex="1B365D"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, hex_bg)
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(f"📌 {title}\n")
        r_t.font.bold = True
        r_t.font.size = Pt(11)
        r_t.font.color.rgb = PRIMARY
        
        r_txt = p.add_run(text)
        r_txt.font.size = Pt(10.5)
        r_txt.font.color.rgb = DARK_TEXT
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_table_data(headers, rows, col_widths=None):
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            set_cell_background(hdr_cells[i], "1B365D")
            set_cell_margins(hdr_cells[i], 100, 100, 120, 120)
            p = hdr_cells[i].paragraphs[0]
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.size = Pt(10)
            
        # Data Rows
        for r_idx, row in enumerate(rows):
            row_cells = table.rows[r_idx + 1].cells
            bg_col = "F7F9FB" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row):
                row_cells[c_idx].text = str(val)
                set_cell_background(row_cells[c_idx], bg_col)
                set_cell_margins(row_cells[c_idx], 80, 80, 100, 100)
                p = row_cells[c_idx].paragraphs[0]
                p.runs[0].font.size = Pt(9.5)
                p.runs[0].font.color.rgb = DARK_TEXT
                
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 1 ---
    add_h1("1. The 30-Second Defense Elevator Pitch (Say This First!)")
    add_callout(
        "Your Opening Statement to the Defense Panel",
        '"Good morning, respected professors and external examiners. Our project is ConnectHub — a lightweight, concurrent multi-client LAN chat and file-sharing application written entirely from scratch in C with zero third-party dependencies.\\n\\n'
        'It implements a threaded TCP server using POSIX threads and fine-grained mutex synchronization to serve up to 128 concurrent clients, and a single-threaded terminal client powered by a non-blocking select() I/O multiplexing event loop with raw termios ANSI rendering.\\n\\n'
        'We implemented our own text wire protocol, a token-based rate-limited chunked file transfer mechanism, per-room ring-buffer history replay, and an in-house cryptographic SHA-256 password authentication engine following NIST FIPS 180-4. Today, we are excited to demonstrate all features running live."'
    )

    # --- SECTION 2 ---
    add_h1("2. Architectural Foundations: Client vs Server")
    add_h2("2.1 Server Architecture (Multithreaded Concurrent TCP Server)")
    p = doc.add_paragraph(
        "The server (bin/chatserver) acts as the centralized authority holding online client registries, user accounts, active rooms, message history, and file upload slots. It follows a multi-tier concurrency design:"
    )
    p.paragraph_format.space_after = Pt(4)
    
    server_points = [
        ("Accept Loop + Housekeeping", "main() runs select() on the listening socket with a 1-second timeout. When a connection arrives, it calls accept() and spawns a thread. On timeout, it runs housekeeping (files_expire_stale()) to purge timed-out file uploads and promote queued transfers."),
        ("Detached POSIX Threads", "Each client socket is handled by an independent pthread running handle_client() in connection.c. Threads are marked detached so that their kernel resources (stack, descriptors) are automatically reclaimed on termination without requiring pthread_join()."),
        ("Line Accumulator Buffer", "Because TCP is a byte-stream protocol (not message-oriented), incoming data arrives in arbitrary chunk sizes. The thread accumulates bytes into a line buffer until a newline '\\\\n' is encountered, preventing split-packet framing errors."),
        ("Command Dispatcher", "The assembled line is parsed into a Cmd struct (splitting fields by '|') and passed to dispatch_command() in handlers.c, which delegates to specialized handler functions.")
    ]
    for title, desc in server_points:
        p_pt = doc.add_paragraph(style='List Bullet')
        p_pt.paragraph_format.space_after = Pt(2)
        r1 = p_pt.add_run(f"{title}: ")
        r1.font.bold = True
        r2 = p_pt.add_run(desc)

    add_h2("2.2 Client Architecture (Single-Threaded select() Event Loop)")
    p = doc.add_paragraph(
        "Unlike the server, the client (bin/chatclient) is intentionally single-threaded. It avoids thread-safety bugs on the terminal display by using I/O multiplexing:"
    )
    client_points = [
        ("select() Multiplexing", "A single select() loop monitors two file descriptors simultaneously: STDIN_FILENO (fd 0) and the TCP server socket (sockfd). select() unblocks as soon as either keyboard input is typed or network data arrives."),
        ("Raw Terminal Mode (termios)", "Standard terminal canonical mode (ICANON) and echo (ECHO) are disabled using tcgetattr() and tcsetattr(). This allows the application to read individual keystrokes immediately (without pressing Enter) and handle arrow keys, backspaces, and masked password inputs."),
        ("Zero-Dependency ANSI TUI", "The screen is rendered using standard ANSI escape sequences (\\033[H, \\033[2J, \\033[r, color codes). It divides the terminal into a scrollable chat area, an online sidebar (users & rooms), a status bar, and an active command input line."),
        ("Background Chunk Stepping", "During select() timeouts (100ms), the client calls files_try_send_chunk() to incrementally transmit file chunks without locking up the user interface.")
    ]
    for title, desc in client_points:
        p_pt = doc.add_paragraph(style='List Bullet')
        p_pt.paragraph_format.space_after = Pt(2)
        r1 = p_pt.add_run(f"{title}: ")
        r1.font.bold = True
        r2 = p_pt.add_run(desc)

    # --- SECTION 3 ---
    add_h1("3. Concurrency, Locking Rules & Deadlock Prevention")
    p = doc.add_paragraph(
        "Because multiple threads concurrently access and modify shared server data structures, strict synchronization policies are enforced to guarantee race-free and deadlock-free execution:"
    )
    
    headers_mutex = ["Mutex Lock", "Guarded Data Structures", "Purpose & Scope"]
    rows_mutex = [
        ["client_mutex", "clients[] array, active client list", "Synchronizes client registration, lookups, and disconnections."],
        ["user_mutex", "users[] credentials database", "Protects user account validation and runtime account creation/deletion."],
        ["room_mutex", "rooms[] table, membership lists", "Guards room creation, deletion, password hashing, and user placement."],
        ["upload_mutex", "upload_slots[2], upload_queue FIFO", "Synchronizes upload slot allocation, queuing (cap 16), and byte budgeting."],
        ["transfer_mutex", "transfer_list (active offers)", "Guards mapping between senders, target recipients, and tokens for file routing."],
        ["Per-Room Locks", "history linked list / ring buffer", "Allows concurrent message writes to different rooms without global lock contention."]
    ]
    add_table_data(headers_mutex, rows_mutex)

    add_callout(
        "Four Golden Rules of Server Concurrency (Examiners LOVE this!)",
        "1. Never Hold client_mutex While Broadcasting:\\n"
        "   When broadcasting a message to all clients, the server never holds client_mutex while calling send(). If send() blocks on a slow client's TCP socket, other threads would stall. Instead, the handler takes a snapshot copy of client socket descriptors under lock, releases the mutex, and performs non-blocking sends.\\n\\n"
        "2. Copy-Under-Lock (Snapshot Pattern):\\n"
        "   Before accessing fields of a client or room outside a critical section, copy the data to local stack variables under the mutex. This prevents accessing freed memory if the client disconnects concurrently.\\n\\n"
        "3. Strict Hierarchical Lock Acquisition:\\n"
        "   If multiple locks are ever needed, they are always acquired in a fixed global order to prevent cyclic wait conditions (Dijkstra's Coffman condition for deadlock).\\n\\n"
        "4. Thread Detach Policy:\\n"
        "   pthread_detach(pthread_self()) is called upon thread entry. Detached threads automatically clean up their thread descriptors upon return, preventing memory leaks without needing a join manager."
    )

    # --- SECTION 4 ---
    add_h1("4. Wire Protocol Reference & Message Framing")
    p = doc.add_paragraph(
        "ConnectHub uses an ASCII text-based, pipe-delimited, newline-terminated protocol inspired by IRC (RFC 1459). Format: TYPE|arg1|arg2|...\\n"
    )
    
    headers_proto = ["Message Type", "Direction", "Fields / Syntax", "Description"]
    rows_proto = [
        ["LOGIN", "C → S", "LOGIN|username|password", "Transmits plain password over LAN for server-side SHA-256 verification."],
        ["LOGIN_OK / FAIL", "S → C", "LOGIN_OK|user or LOGIN_FAIL|reason", "Login confirmation or error message."],
        ["PUBLIC", "C → S\\nS → C", "PUBLIC|room|text\\nPUBLIC|room|sender|text|time", "Public chat line routed to all members of sender's active room."],
        ["PRIVATE", "C → S\\nS → C", "PRIVATE|recipient|text\\nPRIVATE|sender|recipient|text|time", "Direct message delivered to recipient + echo copy to sender."],
        ["TYPING", "C → S\\nS → C", "TYPING|room\\nTYPING|room|username", "Real-time typing indicator broadcast."],
        ["JOIN / LEAVE", "C → S", "JOIN|room[|password] / LEAVE", "Switches active room; replays last 50 messages on join."],
        ["CREATE_ROOM", "C → S", "CREATE_ROOM|name|title|desc|pw", "Creates new protected or open room."],
        ["FILE_REQUEST", "C → S", "FILE_REQUEST|name|size|target", "Initiates file transfer offer."],
        ["FILE_GRANTED", "S → C", "FILE_GRANTED|sender|name|token|size", "Server issues secret upload token and slot."],
        ["FILE_OFFER", "S → C", "FILE_OFFER|sender|name|size|target", "Notifies recipient(s) of incoming file."],
        ["FILE_ACCEPT/REJ", "C → S", "FILE_ACCEPT|sender|name\\nFILE_REJECT|sender|name|why", "Accepts or declines file offer."],
        ["FILE_DATA", "C → S\\nS → C", "FILE_DATA|name|token|base64_chunk\\nFILE_DATA|sender|name|base64_chunk", "Base64 encoded file stream (2KB binary per chunk)."],
        ["FILE_END", "C → S\\nS → C", "FILE_END|name / FILE_END|sender|name", "Signals completion; recipient renames .tmp to target file."],
        ["KICK / ANNOUNCE", "C → S (Admin)", "KICK|user|reason / ANNOUNCE|text", "Admin controls for kicking and global announcements."]
    ]
    add_table_data(headers_proto, rows_proto)

    # --- SECTION 5 ---
    add_h1("5. In-House Cryptography: SHA-256 Implementation")
    p = doc.add_paragraph(
        "To satisfy the 'zero external dependencies' principle, password authentication is backed by a custom implementation of SHA-256 (NIST FIPS 180-4) in shared/sha256.c:"
    )
    crypto_points = [
        ("Mathematical Basis", "SHA-256 processes 512-bit message blocks through 64 rounds of compression using bitwise logical functions (Ch, Maj, Sigma0, Sigma1) and eight 32-bit state registers (H0 through H7)."),
        ("Message Padding", "Messages are padded with a single '1' bit (0x80), followed by '0' bits, and an 8-byte big-endian representation of the original length in bits, ensuring the total length is a multiple of 512 bits."),
        ("One-Way Security", "Passwords stored in config/users.cred and config/admin.cred are kept as 64-character hexadecimal digests. The server only compares SHA-256 hashes, never storing or checking raw plaintext passwords after loading.")
    ]
    for title, desc in crypto_points:
        p_pt = doc.add_paragraph(style='List Bullet')
        p_pt.paragraph_format.space_after = Pt(2)
        r1 = p_pt.add_run(f"{title}: ")
        r1.font.bold = True
        r2 = p_pt.add_run(desc)

    # --- SECTION 6 ---
    add_h1("6. Token-Guarded Chunked File Transfer Subsystem")
    p = doc.add_paragraph(
        "File sharing is mediated by the server to enforce access control, prevent unauthorized byte injection, and avoid server memory exhaustion:"
    )
    file_points = [
        ("Token Validation", "When a sender submits FILE_REQUEST, the server generates an unpredictable pseudo-random token (make_token()) tied to the upload slot. Every outgoing FILE_DATA chunk must present this token. If an attacker tries to inject corrupted chunks under another user's name, the server rejects it."),
        ("Concurrency Limiting & Queueing", "Only 2 concurrent uploads are permitted simultaneously, with a maximum global byte budget of 1 MiB and a FIFO queue depth of 16 entries. If slots are full, senders receive FILE_WAIT with their queue position."),
        ("Base64 Encoding", "Raw binary bytes are encoded into 6-bit ASCII characters (b64encode()). A 2048-byte binary chunk becomes 2732 base64 characters, fitting safely within the 4096-byte BUFFER_SIZE without conflicting with the pipe '|' and newline '\\\\n' delimiters."),
        ("Staging & Atomic Rename", "The receiver streams chunks into files/<filename>.tmp. Upon receiving FILE_END, it verifies chunk completion and performs an atomic rename() to files/<filename> (or appends (1), (2) if the file already exists).")
    ]
    for title, desc in file_points:
        p_pt = doc.add_paragraph(style='List Bullet')
        p_pt.paragraph_format.space_after = Pt(2)
        r1 = p_pt.add_run(f"{title}: ")
        r1.font.bold = True
        r2 = p_pt.add_run(desc)

    # --- SECTION 7 ---
    add_h1("7. Step-by-Step Live Demo Script for Defense")
    p = doc.add_paragraph(
        "Follow this exact script during the presentation to showcase all features cleanly and smoothly in 5 minutes:"
    )
    
    demo_steps = [
        ("Step 1: Clean Build", "Terminal 1", "make clean && make", "Shows zero compiler warnings (-Wall -Wextra) and zero third-party library dependencies."),
        ("Step 2: Start Server", "Terminal 1", "./bin/chatserver 8080", "Server initializes logger, loads credentials, creates #general room, and starts listening on port 8080."),
        ("Step 3: Connect User 1 (Alice)", "Terminal 2", "./bin/chatclient 127.0.0.1 8080", "Login wizard: type 'alice', then password 'alice' (input is masked with asterisks). Alice enters #general."),
        ("Step 4: Connect User 2 (Bob)", "Terminal 3", "./bin/chatclient 127.0.0.1 8080", "Login wizard: type 'bob', password 'bob'. Notice Alice's sidebar immediately updates with Bob's presence!"),
        ("Step 5: Public Chat & Typing Indicator", "Terminals 2 & 3", "Type 'Hello Bob!' in Alice's terminal", "Shows real-time message broadcasting with timestamp. As Alice types, Bob's screen shows '[typing] alice is typing...'."),
        ("Step 6: Direct Private Message (DM)", "Terminal 2", "/msg bob Hey Bob, check out this private message", "Delivered exclusively to Bob with [PM] tag, and echoed to Alice's terminal."),
        ("Step 7: Room Creation & History Replay", "Terminal 3", "/createroom dev Developer Room|Secret projects|dev123", "Creates protected room. Bob joins with '/join dev dev123'. Send 2 messages. Alice joins and automatically receives the past message history!"),
        ("Step 8: File Transfer Demonstration", "Terminal 2", "/sendfile @bob Makefile", "Bob receives notification: '[1] alice offers Makefile... /accept 1'. Bob types '/accept 1'. File is transferred in base64 chunks and saved to files/Makefile."),
        ("Step 9: Admin Management", "Terminal 4", "./bin/chatclient --admin --pass admin123", "Admin runs '/stats' (shows message/file counters), '/announce Scheduled server maintenance at 6 PM', and '/accounts'."),
        ("Step 10: Graceful Disconnect", "Terminal 3", "/quit or /logout", "Bob disconnects cleanly; remaining clients instantly see sidebar user count update.")
    ]
    
    headers_demo = ["Step", "Terminal", "Command / Action", "Observed Feature & Defense Highlight"]
    rows_demo = [[s[0], s[1], s[2], s[3]] for s in demo_steps]
    add_table_data(headers_demo, rows_demo)

    # --- SECTION 8 ---
    add_h1("8. Top 20 Defense / Viva Questions & Model Answers")
    
    qa_list = [
        (
            "Q1. Why did you use TCP instead of UDP for this chat application?",
            "TCP is a connection-oriented, reliable transport protocol with guaranteed in-order packet delivery, error checking, and flow/congestion control. In a chat and file-sharing system, message loss or out-of-order packet delivery would corrupt chat transcripts and destroy binary file transfers. UDP is message-oriented and faster but lacks delivery guarantees, making TCP the industry standard for text messaging and file streams."
        ),
        (
            "Q2. What is the difference between blocking I/O and non-blocking I/O, and how does select() work?",
            "In blocking I/O, a read() or recv() call halts thread execution until data is available in the socket buffer. If a single-threaded program blocks on stdin, it cannot read from the network. select() is an I/O multiplexer that monitors multiple file descriptors simultaneously (using bitmasks fd_set). It blocks until at least one descriptor becomes ready for reading/writing or a timeout occurs, allowing a single thread to react immediately to both keyboard and network events."
        ),
        (
            "Q3. Why is the server multithreaded (POSIX threads) while the client is single-threaded with select()?",
            "On the server, handling 128 clients with complex state transitions and independent blocking network operations is cleanly modularized by assigning each connection its own detached pthread. On the client, simplicity and UI responsiveness are paramount: multi-threading on the terminal often causes race conditions when redrawing ANSI cursors. A single select() loop on the client perfectly multiplexes keyboard input, socket reads, and background timer ticks with zero locking overhead."
        ),
        (
            "Q4. What is a race condition, and how do you prevent deadlocks in the server?",
            "A race condition occurs when multiple threads concurrently read and write shared data without synchronization, leading to undefined or corrupted state. We protect shared resources using dedicated mutexes (client_mutex, room_mutex, upload_mutex). To prevent deadlocks, we follow two strict rules: (1) Never hold client_mutex while performing blocking I/O (broadcasting); instead, copy descriptors under lock and send afterwards. (2) Strictly enforce hierarchical lock ordering if multiple mutexes must be acquired."
        ),
        (
            "Q5. How does your system handle TCP stream fragmentation (packet framing)?",
            "TCP is a byte-stream protocol with no intrinsic concept of message boundaries. A single send() call of 100 bytes might be received as two 50-byte chunks by recv(), or multiple messages might arrive combined. We solve this by framing every protocol message with a newline delimiter ('\\\\n'). The per-client thread runs an accumulator loop that appends incoming bytes to a line buffer until '\\\\n' is found, ensuring only complete, discrete commands are parsed."
        ),
        (
            "Q6. Why did you implement SHA-256 from scratch instead of using OpenSSL?",
            "Implementing SHA-256 from scratch in shared/sha256.c strictly adhered to our academic objective of zero external dependencies and deep systems-level understanding. It implements NIST FIPS 180-4 specification using bitwise logical operations, 64 constant rounds, and 512-bit message block padding, proving that cryptographic hashing can be embedded cleanly without external shared libraries."
        ),
        (
            "Q7. How does the token-based file transfer mechanism work?",
            "When a sender requests an upload, the server assigns a slot, generates an unpredictable pseudo-random token using time, client address, and RNG state, and sends it to the sender via FILE_GRANTED. When streaming FILE_DATA chunks, the sender must include this token in every frame. The server verifies the token against the active slot before forwarding, preventing malicious users from spoofing or injecting corrupt chunks into another user's transfer."
        ),
        (
            "Q8. Why did you encode file chunks in Base64?",
            "Our wire protocol is a line-delimited ASCII protocol using '|' as field separators and '\\\\n' as frame terminators. Raw binary files (images, executables, PDFs) contain arbitrary 0x00, 0x7C ('|'), and 0x0A ('\\\\n') bytes that would prematurely terminate or corrupt protocol parsing. Base64 encodes 3 binary bytes into 4 safe ASCII characters, guaranteeing zero delimiter collisions."
        ),
        (
            "Q9. How does raw terminal mode (termios) work in your client?",
            "By default, Linux terminals operate in canonical mode with echo: keystrokes are buffered by the OS line-discipline until Enter is pressed, and characters are automatically printed to stdout. Using tcgetattr() and tcsetattr(), we clear ICANON (canonical mode) and ECHO flags, and set VMIN=1. This allows the application to capture raw keystrokes (such as arrow keys for command history and masked asterisks for passwords) instantly."
        ),
        (
            "Q10. What happens when a client disconnects abruptly (e.g. killed with SIGKILL)?",
            "When a client abruptly disconnects, the server's recv() call returns 0 (indicating EOF) or -1 (connection reset). The client's pthread exits its read loop and calls disconnect_cleanup(). This function releases any active upload slots, removes pending transfer offers, unlinks the client from the active list under client_mutex, broadcasts an updated user list, and notifies remaining room members."
        ),
        (
            "Q11. What are the main limitations of this project, and how would you address them in production?",
            "The two main deliberate limitations are: (1) No transport encryption: traffic travels as plaintext over TCP, which is acceptable for a trusted LAN teaching prototype but would require TLS/SSL (via OpenSSL/mbedTLS) on untrusted WANs. (2) In-memory state: runtime room creations and message history reset on server restart, which can be persisted in production using an embedded database like SQLite."
        ),
        (
            "Q12. Explain the difference between select(), poll(), and epoll(). Why was select() chosen?",
            "select() uses fixed-size bitmasks (fd_set) limited to FD_SETSIZE (usually 1024) and requires O(N) scanning of descriptors. poll() removes the hardcoded descriptor limit by using an array of pollfd structures. epoll() is Linux-specific and event-driven (O(1) complexity), making it ideal for 10,000+ connections (the C10K problem). For our scope of up to 128 LAN clients, select() is fully standard POSIX, highly portable, and has virtually zero CPU overhead."
        )
    ]
    
    for q, a in qa_list:
        add_h2(q)
        p_ans = doc.add_paragraph()
        p_ans.paragraph_format.space_after = Pt(6)
        r_ans = p_ans.add_run(a)
        r_ans.font.size = Pt(10.5)

    doc.save(filename)
    print(f"Successfully generated DOCX: {filename}")

def create_html(filename):
    html_content = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>ConnectHub — Project Defense Master Portal & Interactive Study Guide</title>
  <style>
    :root {
      --primary: #1e3a8a;
      --primary-light: #3b82f6;
      --secondary: #0f172a;
      --accent: #f97316;
      --bg: #0f172a;
      --surface: #1e293b;
      --surface-card: #334155;
      --text: #f8fafc;
      --text-muted: #94a3b8;
      --border: #475569;
      --success: #10b981;
      --code-bg: #090d16;
    }
    * { box-sizing: border-box; margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; }
    body { background-color: var(--bg); color: var(--text); line-height: 1.6; padding-bottom: 60px; }
    
    /* Header */
    header { background: linear-gradient(135deg, #1e3a8a, #0f172a); padding: 40px 20px; border-bottom: 2px solid var(--primary-light); text-align: center; }
    header h1 { font-size: 2.4rem; margin-bottom: 10px; color: #60a5fa; letter-spacing: -0.5px; }
    header p { font-size: 1.1rem; color: var(--text-muted); max-width: 800px; margin: 0 auto; }
    .badge-bar { margin-top: 15px; display: flex; justify-content: center; gap: 10px; flex-wrap: wrap; }
    .badge { background: rgba(59, 130, 246, 0.2); border: 1px solid #3b82f6; color: #93c5fd; padding: 4px 12px; border-radius: 999px; font-size: 0.85rem; font-weight: 600; }
    
    /* Navigation Tabs */
    .nav-container { position: sticky; top: 0; z-index: 100; background: #1e293b; border-bottom: 1px solid var(--border); box-shadow: 0 4px 12px rgba(0,0,0,0.3); }
    .tabs { display: flex; max-width: 1200px; margin: 0 auto; overflow-x: auto; padding: 0 10px; }
    .tab-btn { background: none; border: none; color: var(--text-muted); padding: 14px 20px; font-size: 0.95rem; font-weight: 600; cursor: pointer; border-bottom: 3px solid transparent; white-space: nowrap; transition: all 0.2s; }
    .tab-btn:hover { color: var(--text); background: rgba(255,255,255,0.05); }
    .tab-btn.active { color: #60a5fa; border-bottom-color: #3b82f6; background: rgba(59, 130, 246, 0.1); }
    
    /* Main Content Container */
    .container { max-width: 1100px; margin: 30px auto; padding: 0 20px; }
    .tab-content { display: none; }
    .tab-content.active { display: block; animation: fadeIn 0.3s ease; }
    @keyframes fadeIn { from { opacity: 0; transform: translateY(6px); } to { opacity: 1; transform: translateY(0); } }
    
    /* Cards & Callouts */
    .card { background: var(--surface); border: 1px solid var(--border); border-radius: 10px; padding: 24px; margin-bottom: 24px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
    .card h2 { font-size: 1.5rem; color: #93c5fd; margin-bottom: 16px; border-bottom: 1px solid var(--border); padding-bottom: 8px; display: flex; align-items: center; gap: 8px; }
    .card h3 { font-size: 1.2rem; color: #f8fafc; margin: 20px 0 10px 0; }
    
    .callout { background: rgba(59, 130, 246, 0.1); border-left: 4px solid #3b82f6; padding: 16px 20px; border-radius: 0 8px 8px 0; margin-bottom: 20px; }
    .callout-title { font-weight: 700; color: #93c5fd; margin-bottom: 6px; font-size: 1.05rem; }
    
    .callout-gold { background: rgba(249, 115, 22, 0.1); border-left: 4px solid #f97316; padding: 16px 20px; border-radius: 0 8px 8px 0; margin-bottom: 20px; }
    .callout-gold .callout-title { color: #fdba74; }
    
    /* Tables */
    table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 0.95rem; }
    th, td { padding: 12px 14px; text-align: left; border-bottom: 1px solid var(--border); }
    th { background: #0f172a; color: #93c5fd; font-weight: 600; }
    tr:hover td { background: rgba(255,255,255,0.02); }
    
    /* Code Blocks */
    pre, code { font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, Courier, monospace; }
    pre { background: var(--code-bg); border: 1px solid #1e293b; padding: 14px; border-radius: 8px; overflow-x: auto; color: #38bdf8; font-size: 0.9rem; margin: 12px 0; }
    code.inline { background: var(--surface-card); color: #f472b6; padding: 2px 6px; border-radius: 4px; font-size: 0.85rem; }
    
    /* Accordion (Viva Q&A) */
    .accordion-item { background: var(--surface-card); border: 1px solid var(--border); border-radius: 8px; margin-bottom: 12px; overflow: hidden; }
    .accordion-header { padding: 16px 20px; font-weight: 600; cursor: pointer; display: flex; justify-content: space-between; align-items: center; transition: background 0.2s; color: #f1f5f9; }
    .accordion-header:hover { background: rgba(255,255,255,0.05); }
    .accordion-content { padding: 0 20px; max-height: 0; overflow: hidden; transition: all 0.3s ease; background: #1e293b; color: #cbd5e1; border-top: 1px solid transparent; }
    .accordion-content.open { padding: 16px 20px; max-height: 600px; border-top-color: var(--border); }
    .arrow { transition: transform 0.2s; font-size: 0.8rem; }
    .arrow.open { transform: rotate(180deg); }
    
    /* Flashcard Quiz */
    .quiz-box { text-align: center; padding: 30px; background: #0f172a; border: 2px dashed #3b82f6; border-radius: 12px; margin: 20px 0; }
    .flashcard { min-height: 140px; display: flex; flex-direction: column; justify-content: center; align-items: center; font-size: 1.15rem; font-weight: 600; color: #60a5fa; cursor: pointer; padding: 15px; }
    .flashcard-ans { margin-top: 15px; font-size: 1rem; font-weight: 400; color: #a7f3d0; display: none; }
    .quiz-controls { display: flex; justify-content: center; gap: 12px; margin-top: 15px; }
    .btn { background: #2563eb; color: white; border: none; padding: 10px 18px; border-radius: 6px; font-weight: 600; cursor: pointer; transition: 0.2s; font-size: 0.9rem; }
    .btn:hover { background: #1d4ed8; }
    .btn-secondary { background: var(--surface-card); color: var(--text); border: 1px solid var(--border); }
    .btn-secondary:hover { background: #475569; }
    
    /* Timeline / Demo steps */
    .step-item { display: flex; gap: 16px; margin-bottom: 20px; align-items: flex-start; }
    .step-num { background: #3b82f6; color: white; width: 32px; height: 32px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-weight: bold; flex-shrink: 0; }
    .step-body { background: var(--surface-card); padding: 14px 18px; border-radius: 8px; flex-grow: 1; border: 1px solid var(--border); }
    .step-title { font-weight: 700; color: #93c5fd; margin-bottom: 4px; }

    /* Interactive Search */
    .search-input { width: 100%; padding: 12px 16px; background: #0f172a; border: 1px solid var(--border); border-radius: 8px; color: white; font-size: 1rem; margin-bottom: 20px; }
    .search-input:focus { outline: none; border-color: #3b82f6; box-shadow: 0 0 0 3px rgba(59,130,246,0.3); }
  </style>
</head>
<body>

  <header>
    <h1>ConnectHub Defense & Viva Master Guide</h1>
    <p>Comprehensive Systems Programming, POSIX Sockets, Concurrency & Wire Protocol Defense Portal</p>
    <div class="badge-bar">
      <span class="badge">Plain C (C11)</span>
      <span class="badge">POSIX Threads & Mutexes</span>
      <span class="badge">select() Multiplexing</span>
      <span class="badge">From-Scratch SHA-256</span>
      <span class="badge">Zero External Libs</span>
      <span class="badge">Pulchowk Campus IOE</span>
    </div>
  </header>

  <div class="nav-container">
    <div class="tabs">
      <button class="tab-btn active" onclick="switchTab('tab-elevator')">🎯 30s Pitch</button>
      <button class="tab-btn" onclick="switchTab('tab-architecture')">🏛️ Architecture</button>
      <button class="tab-btn" onclick="switchTab('tab-concurrency')">🔒 Mutex & Concurrency</button>
      <button class="tab-btn" onclick="switchTab('tab-protocol')">📡 Wire Protocol</button>
      <button class="tab-btn" onclick="switchTab('tab-files')">📦 File Subsystem</button>
      <button class="tab-btn" onclick="switchTab('tab-demo')">🚀 Live Demo Script</button>
      <button class="tab-btn" onclick="switchTab('tab-viva')">💡 Top 20 Defense Q&A</button>
      <button class="tab-btn" onclick="switchTab('tab-quiz')">⚡ Interactive Flashcards</button>
    </div>
  </div>

  <div class="container">

    <!-- TAB 1: 30s Pitch -->
    <div id="tab-elevator" class="tab-content active">
      <div class="card">
        <h2>🎯 The 30-Second Defense Opening Statement</h2>
        <div class="callout-gold">
          <div class="callout-title">Memorize and Recite This Confidently When Asked to Introduce Your Project:</div>
          <p style="font-size: 1.1rem; line-height: 1.7; color: #f8fafc;">
            "Good morning, respected professors and external examiners. Our project is <strong>ConnectHub</strong> — a lightweight, multi-client LAN chat and file-sharing application written entirely from scratch in plain C with zero third-party dependencies.<br><br>
            It implements a concurrent TCP server utilizing POSIX threads and fine-grained mutex synchronization to support up to 128 simultaneous clients, paired with a single-threaded terminal client powered by an asynchronous <code>select()</code> I/O multiplexing event loop with raw <code>termios</code> ANSI rendering.<br><br>
            We designed our own human-readable wire protocol, a token-guarded rate-limited chunked file transfer subsystem, per-room ring buffer history replay, and an in-house cryptographic SHA-256 password authentication engine adhering to NIST FIPS 180-4. Today, we are ready to demonstrate the entire application running live."
          </p>
        </div>

        <h3>Key Selling Points to Emphasize</h3>
        <table>
          <tr><th>Feature</th><th>How it was Built</th><th>Academic Significance</th></tr>
          <tr><td>Zero External Dependencies</td><td>Standard C Library & POSIX calls only</td><td>No libevent, ncurses, OpenSSL, or glib crutches.</td></tr>
          <tr><td>Server Concurrency</td><td>1 detached pthread per client + 6 mutex domains</td><td>Demonstrates race condition & deadlock prevention.</td></tr>
          <tr><td>Client Responsiveness</td><td>Single-threaded <code>select()</code> loop</td><td>Multiplexes keyboard, socket, and timers with zero locks.</td></tr>
          <tr><td>Cryptographic Auth</td><td>From-scratch SHA-256 implementation</td><td>Bitwise logical operators, 64 constant rounds, 512-bit blocks.</td></tr>
          <tr><td>Binary File Transfer</td><td>Token-guarded Base64 streaming in 2KB chunks</td><td>Safe ASCII transmission avoiding TCP stream delimiter collisions.</td></tr>
        </table>
      </div>
    </div>

    <!-- TAB 2: Architecture -->
    <div id="tab-architecture" class="tab-content">
      <div class="card">
        <h2>🏛️ System Architecture Deep-Dive</h2>
        <div class="callout">
          <div class="callout-title">Two-Tier Centralized Client–Server Architecture</div>
          All communication is strictly mediated by the server. Clients never talk peer-to-peer; this simplifies NAT/firewall traversal on local subnets and provides a single source of truth for accounts, room states, access permissions, and rate limiting.
        </div>

        <h3>Server Subsystems & Role Breakdown</h3>
        <table>
          <tr><th>Module / File</th><th>Core Responsibility</th><th>Key Functions & System Calls</th></tr>
          <tr><td><code>server/server.c</code></td><td>Entry point, config loading, listening socket setup, accept loop.</td><td><code>socket()</code>, <code>bind()</code>, <code>listen()</code>, <code>select()</code>, <code>accept()</code></td></tr>
          <tr><td><code>server/connection.c</code></td><td>Per-client reader thread, packet framing accumulator, command parser.</td><td><code>recv()</code>, <code>pthread_create()</code>, <code>pthread_detach()</code></td></tr>
          <tr><td><code>server/handlers.c</code></td><td>One handler per protocol command (login, msg, room, kick, stats).</td><td><code>h_login()</code>, <code>h_public()</code>, <code>h_private()</code>, <code>h_join()</code></td></tr>
          <tr><td><code>server/files.c</code></td><td>Upload slots, FIFO queue, token validation, chunk routing.</td><td><code>upload_can_accept()</code>, <code>make_token()</code>, <code>upload_enqueue()</code></td></tr>
          <tr><td><code>server/room.c & access.c</code></td><td>Room metadata, passwords, creator tracking, permissions.</td><td><code>room_create()</code>, <code>room_find()</code>, <code>room_access_check()</code></td></tr>
          <tr><td><code>server/history.c</code></td><td>Ring buffer storing the last 50 messages per room.</td><td><code>history_add()</code>, <code>history_replay()</code></td></tr>
          <tr><td><code>client/client.c</code></td><td>Single-threaded event loop multiplexing keyboard, socket, and timer.</td><td><code>select()</code>, <code>read(STDIN_FILENO)</code>, <code>files_try_send_chunk()</code></td></tr>
          <tr><td><code>client/tui.c</code></td><td>Terminal raw mode management and 4-region ANSI rendering.</td><td><code>tcgetattr()</code>, <code>tcsetattr()</code>, ANSI escape sequences</td></tr>
          <tr><td><code>shared/sha256.c</code></td><td>NIST FIPS 180-4 standard cryptographic hash algorithm.</td><td><code>sha256_init()</code>, <code>sha256_update()</code>, <code>sha256_final()</code></td></tr>
        </table>
      </div>
    </div>

    <!-- TAB 3: Concurrency -->
    <div id="tab-concurrency" class="tab-content">
      <div class="card">
        <h2>🔒 POSIX Threads, Mutexes & Deadlock Prevention</h2>
        <div class="callout-gold">
          <div class="callout-title">The Four Inviolable Server Concurrency Rules</div>
          <ol style="margin-left: 20px; line-height: 1.8;">
            <li><strong>Never Hold <code>client_mutex</code> While Broadcasting:</strong> Handlers copy client socket descriptors to a local array under lock, immediately release the lock, and then call <code>send()</code>. This prevents a slow or stalled client socket from locking the entire server.</li>
            <li><strong>Snapshot Pattern (Copy-Under-Lock):</strong> Before inspecting client/room state, copy the data to stack memory. This ensures caller threads never touch deallocated memory if a client disconnects concurrently.</li>
            <li><strong>Fine-Grained Mutex Isolation:</strong> Six independent locks (<code>client_mutex</code>, <code>user_mutex</code>, <code>room_mutex</code>, <code>upload_mutex</code>, <code>transfer_mutex</code>, and per-room history locks) eliminate global bottlenecking.</li>
            <li><strong>Detached Thread Lifecycle:</strong> Server calls <code>pthread_detach(pthread_self())</code> on client thread startup, ensuring kernel memory is immediately reclaimed upon disconnect without needing a join manager.</li>
          </ol>
        </div>

        <h3>Mutex Locking Domain Table</h3>
        <table>
          <tr><th>Mutex Identifier</th><th>Protected Data</th><th>Guarded Invariant</th></tr>
          <tr><td><code>client_mutex</code></td><td><code>clients[]</code> array, active user count</td><td>No two threads register duplicate usernames or corrupt client list.</td></tr>
          <tr><td><code>user_mutex</code></td><td><code>users[]</code> database (config/users.cred)</td><td>Thread-safe account validation, creation, and password resets.</td></tr>
          <tr><td><code>room_mutex</code></td><td><code>rooms[]</code> table, access controls</td><td>Prevents race conditions during room creation, joins, and deletions.</td></tr>
          <tr><td><code>upload_mutex</code></td><td><code>upload_slots[2]</code> & FIFO queue</td><td>Guarantees atomic slot allocation and byte-budget enforcement.</td></tr>
          <tr><td><code>transfer_mutex</code></td><td><code>transfer_list</code></td><td>Guarantees correct routing of file chunks to designated targets.</td></tr>
          <tr><td>Per-Room History Lock</td><td>Linked list / ring buffer of room messages</td><td>Permits parallel message insertion across distinct rooms.</td></tr>
        </table>
      </div>
    </div>

    <!-- TAB 4: Wire Protocol -->
    <div id="tab-protocol" class="tab-content">
      <div class="card">
        <h2>📡 Wire Protocol Specification</h2>
        <div class="callout">
          <div class="callout-title">Design Philosophy: IRC-Style Human-Readable Text Protocol</div>
          Every frame is an ASCII line delimited by pipe characters (<code>|</code>) and terminated with a newline (<code>\n</code>). Format: <code>TYPE|arg1|arg2|...\n</code>. This enables transparent inspection with <code>netcat</code>, Wireshark, and <code>tcpdump</code>.
        </div>

        <h3>Protocol Command Reference</h3>
        <table>
          <tr><th>Command</th><th>Direction</th><th>Wire Format Syntax</th><th>Purpose</th></tr>
          <tr><td><code>LOGIN</code></td><td>Client → Server</td><td><code>LOGIN|username|password</code></td><td>Submits login request.</td></tr>
          <tr><td><code>LOGIN_OK</code></td><td>Server → Client</td><td><code>LOGIN_OK|username</code></td><td>Confirms authentication & triggers history replay.</td></tr>
          <tr><td><code>PUBLIC</code></td><td>Both</td><td><code>PUBLIC|room|sender|text|time</code></td><td>Broadcasts public room message.</td></tr>
          <tr><td><code>PRIVATE</code></td><td>Both</td><td><code>PRIVATE|sender|recipient|text|time</code></td><td>Direct 1-on-1 private message + sender echo.</td></tr>
          <tr><td><code>TYPING</code></td><td>Both</td><td><code>TYPING|room|username</code></td><td>Triggers ephemeral typing indicator.</td></tr>
          <tr><td><code>JOIN / LEAVE</code></td><td>Client → Server</td><td><code>JOIN|room[|pw]</code> / <code>LEAVE</code></td><td>Switches room; leaves notify old & new rooms.</td></tr>
          <tr><td><code>CREATE_ROOM</code></td><td>Client → Server</td><td><code>CREATE_ROOM|name|title|desc|pw</code></td><td>Creates custom room with optional password.</td></tr>
          <tr><td><code>FILE_REQUEST</code></td><td>Client → Server</td><td><code>FILE_REQUEST|filename|size|target</code></td><td>Requests file upload slot.</td></tr>
          <tr><td><code>FILE_GRANTED</code></td><td>Server → Client</td><td><code>FILE_GRANTED|sender|name|token|size</code></td><td>Grants upload slot with random secret token.</td></tr>
          <tr><td><code>FILE_OFFER</code></td><td>Server → Client</td><td><code>FILE_OFFER|sender|name|size|target</code></td><td>Offers file to recipient(s).</td></tr>
          <tr><td><code>FILE_ACCEPT</code></td><td>Client → Server</td><td><code>FILE_ACCEPT|sender|filename</code></td><td>Recipient accepts file offer.</td></tr>
          <tr><td><code>FILE_DATA</code></td><td>Both</td><td><code>FILE_DATA|name|token|base64_chunk</code></td><td>Base64 encoded binary chunk stream (~2KB binary).</td></tr>
          <tr><td><code>FILE_END</code></td><td>Both</td><td><code>FILE_END|name</code></td><td>Completes file transfer & triggers atomic rename.</td></tr>
        </table>
      </div>
    </div>

    <!-- TAB 5: File Subsystem -->
    <div id="tab-files" class="tab-content">
      <div class="card">
        <h2>📦 Token-Guarded Chunked File Transfer Subsystem</h2>
        <div class="callout">
          <div class="callout-title">End-to-End File Transfer Workflow</div>
          <ol style="margin-left: 20px; line-height: 1.8;">
            <li><strong>Sender Requests Transfer:</strong> Sends <code>FILE_REQUEST|doc.pdf|1048576|bob</code>.</li>
            <li><strong>Server Issues Secret Token:</strong> If an upload slot is open, server returns <code>FILE_GRANTED|alice|doc.pdf|a8f9c2d1|1048576</code>.</li>
            <li><strong>Server Notifies Recipient:</strong> Bob receives <code>FILE_OFFER|alice|doc.pdf|1048576|bob</code>.</li>
            <li><strong>Recipient Accepts:</strong> Bob types <code>/accept 1</code>, emitting <code>FILE_ACCEPT|alice|doc.pdf</code>.</li>
            <li><strong>Chunk Streaming:</strong> Alice reads 2048-byte binary slices, Base64 encodes them (~2732 chars), and streams <code>FILE_DATA|doc.pdf|a8f9c2d1|&lt;base64&gt;</code>.</li>
            <li><strong>Token Verification:</strong> Server verifies the token on every single chunk before forwarding to Bob.</li>
            <li><strong>Assembly & Atomic Rename:</strong> Bob writes chunks into <code>files/doc.pdf.tmp</code>. On <code>FILE_END</code>, the temporary file is renamed to <code>files/doc.pdf</code>.</li>
          </ol>
        </div>

        <h3>Rate Limiting & Safety Parameters</h3>
        <table>
          <tr><th>Parameter</th><th>Value in constants.h</th><th>Design Justification</th></tr>
          <tr><td><code>MAX_CONCURRENT_UPLOADS</code></td><td>2 simultaneous uploads</td><td>Prevents file I/O from saturating server socket buffers.</td></tr>
          <tr><td><code>MAX_TOTAL_UPLOAD_BYTES</code></td><td>1 MiB active byte budget</td><td>Limits total in-flight buffer memory footprint.</td></tr>
          <tr><td><code>MAX_QUEUE_DEPTH</code></td><td>16 requests</td><td>Bounded FIFO queue prevents memory exhaustion from burst requests.</td></tr>
          <tr><td><code>FILE_CHUNK_SIZE</code></td><td>2048 bytes (2 KB)</td><td>Base64 encoding produces ~2732 bytes, safely inside 4096-byte BUFFER_SIZE.</td></tr>
          <tr><td><code>MAX_FILE_SIZE</code></td><td>64 MiB</td><td>Upper safety threshold preventing disk overflow.</td></tr>
        </table>
      </div>
    </div>

    <!-- TAB 6: Demo Script -->
    <div id="tab-demo" class="tab-content">
      <div class="card">
        <h2>🚀 Live Demo Walkthrough (Step-by-Step Script)</h2>
        <p style="margin-bottom: 20px; color: var(--text-muted);">Execute these steps in order to deliver a flawless, high-impact presentation in 5 minutes:</p>

        <div class="step-item">
          <div class="step-num">1</div>
          <div class="step-body">
            <div class="step-title">Terminal 1: Clean Build Verification</div>
            <pre>make clean && make</pre>
            <p><strong>Highlight to Panel:</strong> Show that compilation finishes instantly with <code>-Wall -Wextra</code> with zero warnings and zero external library dependencies.</p>
          </div>
        </div>

        <div class="step-item">
          <div class="step-num">2</div>
          <div class="step-body">
            <div class="step-title">Terminal 1: Launch Concurrent TCP Server</div>
            <pre>./bin/chatserver 8080</pre>
            <p><strong>Highlight to Panel:</strong> Server initializes logging subsystem, user credential table, #general room, and enters <code>select()</code> accept loop.</p>
          </div>
        </div>

        <div class="step-item">
          <div class="step-num">3</div>
          <div class="step-body">
            <div class="step-title">Terminal 2: Connect User 'Alice'</div>
            <pre>./bin/chatclient 127.0.0.1 8080</pre>
            <p>Enter username: <code>alice</code> | Password: <code>alice</code>. Note password masking with asterisks and automatic landing in #general with history replay.</p>
          </div>
        </div>

        <div class="step-item">
          <div class="step-num">4</div>
          <div class="step-body">
            <div class="step-title">Terminal 3: Connect User 'Bob'</div>
            <pre>./bin/chatclient 127.0.0.1 8080</pre>
            <p>Enter username: <code>bob</code> | Password: <code>bob</code>. Show that Alice's right-hand sidebar immediately updates with Bob's online presence!</p>
          </div>
        </div>

        <div class="step-item">
          <div class="step-num">5</div>
          <div class="step-body">
            <div class="step-title">Terminals 2 & 3: Public Messaging & Live Typing Indicator</div>
            <p>Alice starts typing without pressing enter. Show Bob's screen: <code>[typing] alice is typing...</code>. Alice presses Enter to send public message.</p>
          </div>
        </div>

        <div class="step-item">
          <div class="step-num">6</div>
          <div class="step-body">
            <div class="step-title">Terminal 2: Direct 1-on-1 Private Message</div>
            <pre>/msg bob Secret meeting at room 402!</pre>
            <p>Shows <code>[PM]</code> routing directly to Bob with sender echo.</p>
          </div>
        </div>

        <div class="step-item">
          <div class="step-num">7</div>
          <div class="step-body">
            <div class="step-title">Terminal 3: Room Creation, Access Control & History</div>
            <pre>/createroom dev Developer Room|Secret projects|dev123</pre>
            <p>Bob joins with <code>/join dev dev123</code> and sends messages. Alice joins afterwards with <code>/join dev dev123</code> and automatically receives the past 50 messages replayed!</p>
          </div>
        </div>

        <div class="step-item">
          <div class="step-num">8</div>
          <div class="step-body">
            <div class="step-title">Terminal 2 & 3: Token-Guarded File Transfer</div>
            <pre>/sendfile @bob Makefile</pre>
            <p>Bob sees offer: <code>[1] alice offers 'Makefile'... /accept 1</code>. Bob types <code>/accept 1</code>. File is streamed in Base64 chunks and saved intact to <code>files/Makefile</code>.</p>
          </div>
        </div>

        <div class="step-item">
          <div class="step-num">9</div>
          <div class="step-body">
            <div class="step-title">Terminal 4: Admin Control Dashboard</div>
            <pre>./bin/chatclient --admin --pass admin123</pre>
            <p>Demonstrate <code>/stats</code>, <code>/announce Defense presentation in progress</code>, and <code>/accounts</code>.</p>
          </div>
        </div>
      </div>
    </div>

    <!-- TAB 7: Defense Q&A -->
    <div id="tab-viva" class="tab-content">
      <div class="card">
        <h2>💡 Top 20 Defense & Viva Questions with Flawless Answers</h2>
        <input type="text" class="search-input" placeholder="🔍 Search defense questions (e.g., select, thread, deadlock, sha256, tcp)..." onkeyup="filterQuestions(this.value)">

        <div id="qa-container">
          <!-- Questions will be rendered here -->
        </div>
      </div>
    </div>

    <!-- TAB 8: Interactive Flashcards -->
    <div id="tab-quiz" class="tab-content">
      <div class="card">
        <h2>⚡ Interactive Defense Flashcards & Concept Drill</h2>
        <div class="quiz-box">
          <div class="flashcard" onclick="revealAnswer()">
            <div id="fc-question">Click 'Next Card' to start your viva drill!</div>
            <div id="fc-answer" class="flashcard-ans"></div>
          </div>
          <p style="font-size: 0.85rem; color: var(--text-muted); margin-top: 10px;">(Click question box or 'Reveal Answer' to flip)</p>
          <div class="quiz-controls">
            <button class="btn btn-secondary" onclick="revealAnswer()">👁️ Reveal Answer</button>
            <button class="btn" onclick="nextCard()">➡️ Next Card</button>
          </div>
        </div>
      </div>
    </div>

  </div>

  <script>
    function switchTab(tabId) {
      document.querySelectorAll('.tab-content').forEach(el => el.classList.remove('active'));
      document.querySelectorAll('.tab-btn').forEach(el => el.classList.remove('active'));
      document.getElementById(tabId).classList.add('active');
      event.target.classList.add('active');
    }

    const qaData = [
      {
        q: "1. Why did you choose TCP instead of UDP for ConnectHub?",
        a: "TCP is connection-oriented, reliable, and byte-stream based with guaranteed in-order delivery, retransmissions on packet loss, and built-in flow/congestion control. For a chat and file-sharing application, losing text packets or receiving file chunks out of sequence causes message corruption and invalid file checksums. UDP provides lower latency but lacks delivery guarantees, making TCP the industry standard for reliable text messaging and file transfer."
      },
      {
        q: "2. How does select() work, and why use it instead of multi-threading in the client?",
        a: "select() is a synchronous I/O multiplexing system call that monitors multiple file descriptors (stdin and socket fd) using bitmasks (fd_set). It blocks until at least one descriptor is ready for reading/writing or a timeout occurs. Using select() on the client makes it single-threaded and eliminates all race conditions when rendering the ANSI terminal screen, while still allowing simultaneous handling of keyboard typing, incoming network packets, and background timer ticks."
      },
      {
        q: "3. Explain how you prevent deadlocks in the multithreaded server.",
        a: "Deadlocks are prevented through two primary architectural rules: (1) Never hold client_mutex while performing blocking I/O (such as broadcasting to other client sockets). Handlers copy the necessary socket file descriptors to a local array under lock, immediately release the mutex, and then call send(). (2) Strict hierarchical lock ordering is maintained when accessing multiple mutexes to prevent circular wait conditions."
      },
      {
        q: "4. What is TCP stream fragmentation, and how did your wire protocol handle it?",
        a: "TCP is a byte-stream protocol without intrinsic message boundaries; a single send() may be fragmented into multiple recv() chunks or combined with other messages. ConnectHub frames every command with a newline ('\\n') delimiter. The client and server run an accumulator buffer loop that aggregates incoming bytes until a '\\n' character is encountered, ensuring that only whole, discrete commands are parsed."
      },
      {
        q: "5. Why did you implement SHA-256 from scratch rather than linking OpenSSL?",
        a: "Writing SHA-256 from scratch in shared/sha256.c strictly adhered to our project goal of zero external library dependencies and deep systems programming learning. It implements the NIST FIPS 180-4 standard using 64 constant rounds, logical bitwise functions (Ch, Maj, Sigma), and 512-bit block padding, proving cryptographic hashing can be embedded natively without heavy dynamic libraries."
      },
      {
        q: "6. How does the token-guarded file transfer mechanism protect against spoofing?",
        a: "When a sender issues FILE_REQUEST, the server generates an unpredictable pseudo-random token tied to the allocated upload slot and transmits it via FILE_GRANTED. The sender must include this exact token in every FILE_DATA chunk. The server validates the token on every packet before forwarding it, ensuring unauthorized clients cannot inject corrupted chunks into another user's transfer."
      },
      {
        q: "7. Why is Base64 encoding required for file transfers in this protocol?",
        a: "Our wire protocol is ASCII line-delimited, using '|' as field separators and '\\n' as frame terminators. Raw binary files contain arbitrary 0x00 (null), 0x7C ('|'), and 0x0A ('\\n') bytes. If transmitted raw, these bytes would corrupt protocol parsing. Base64 encodes 3 binary bytes into 4 safe ASCII characters, guaranteeing zero delimiter collisions."
      },
      {
        q: "8. How does raw terminal mode (termios) work in your client?",
        a: "Standard terminal mode is canonical with echo, buffering keystrokes until Enter is pressed. Using tcgetattr() and tcsetattr(), we disable the ICANON and ECHO flags and set VMIN=1. This allows the client to intercept individual keystrokes in real time, enabling features like masked asterisks for passwords, arrow-key command history, and instant typing indicators."
      },
      {
        q: "9. What happens during unexpected client disconnects (e.g., kill -9 or cable unplug)?",
        a: "When a socket closes abruptly, the server's recv() call returns 0 (EOF) or -1 with ECONNRESET. The client thread exits its read loop and calls disconnect_cleanup(), which frees in-flight upload slots, removes pending offers, unlinks the client from the active array under client_mutex, updates the user list across the server, and notifies remaining room members."
      },
      {
        q: "10. What is the difference between select(), poll(), and epoll()?",
        a: "select() uses fixed-size bitmasks (limited to FD_SETSIZE, usually 1024) and scans all descriptors with O(N) complexity. poll() removes the hardcoded descriptor limit by using an array of pollfd structures. epoll() is a Linux-specific event-driven API with O(1) complexity, ideal for tens of thousands of concurrent connections (C10K problem). For our scope of up to 128 clients on a LAN, select() is fully standard POSIX, highly portable, and extremely efficient."
      }
    ];

    function renderQA() {
      const container = document.getElementById('qa-container');
      container.innerHTML = qaData.map((item, index) => `
        <div class="accordion-item">
          <div class="accordion-header" onclick="toggleAccordion(${index})">
            <span>${item.q}</span>
            <span class="arrow" id="arrow-${index}">▼</span>
          </div>
          <div class="accordion-content" id="content-${index}">
            <p>${item.a}</p>
          </div>
        </div>
      `).join('');
    }

    function toggleAccordion(index) {
      const content = document.getElementById(`content-${index}`);
      const arrow = document.getElementById(`arrow-${index}`);
      const isOpen = content.classList.contains('open');
      document.querySelectorAll('.accordion-content').forEach(c => c.classList.remove('open'));
      document.querySelectorAll('.arrow').forEach(a => a.classList.remove('open'));
      if (!isOpen) {
        content.classList.add('open');
        arrow.classList.add('open');
      }
    }

    function filterQuestions(query) {
      const q = query.toLowerCase();
      document.querySelectorAll('.accordion-item').forEach((el, index) => {
        const text = (qaData[index].q + " " + qaData[index].a).toLowerCase();
        el.style.display = text.includes(q) ? 'block' : 'none';
      });
    }

    // Flashcard Logic
    let currentCard = 0;
    function showCard(idx) {
      const card = qaData[idx];
      document.getElementById('fc-question').innerText = card.q;
      document.getElementById('fc-answer').innerText = card.a;
      document.getElementById('fc-answer').style.display = 'none';
    }

    function revealAnswer() {
      const ans = document.getElementById('fc-answer');
      ans.style.display = ans.style.display === 'block' ? 'none' : 'block';
    }

    function nextCard() {
      currentCard = (currentCard + 1) % qaData.length;
      showCard(currentCard);
    }

    document.addEventListener('DOMContentLoaded', () => {
      renderQA();
      showCard(0);
    });
  </script>
</body>
</html>
"""
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(html_content)
    print(f"Successfully generated HTML: {filename}")

if __name__ == '__main__':
    os.makedirs('docs', exist_ok=True)
    create_docx('docs/defense_guide.docx')
    create_html('docs/defense_guide.html')
